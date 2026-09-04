"""
Nepali Legal OCR Correction Script
Uses a selectable OpenAI-compatible API provider.
Includes: Anti-loop penalties, Pre-cleaning, Multi-threading,
          Adaptive chunking, Output validation, and ZWNJ stripping.
"""

import argparse
import datetime
import os
from pathlib import Path
import re
import sys
import time
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Optional, Sequence, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI, RateLimitError
import tiktoken

# ──────────────────────────────────────────────────────────────────
# 1. CONFIGURATION
# ──────────────────────────────────────────────────────────────────
stop_execution = False
SCRIPT_DIR = Path(__file__).resolve().parent
load_dotenv()

ERROR_LOG_PATH = Path(os.getenv("ERROR_LOG_PATH", SCRIPT_DIR / "error_logs.txt"))
CULPRIT_LOG_PATH = Path(os.getenv("CULPRIT_LOG_PATH", SCRIPT_DIR / "skipped_culprits.txt"))

INPUT_DIR_DEFAULT = os.getenv("INPUT_DIR", "data")
OUTPUT_DIR_DEFAULT = os.getenv("OUTPUT_DIR", "output")
DEFAULT_API_PROVIDER = os.getenv("DEFAULT_API_PROVIDER", "deepseek").strip().lower()

# ── Concurrency & Chunk Sizing Defaults ──
MAX_WORKERS_DEFAULT = int(os.getenv("MAX_WORKERS", "5"))
MAX_CHUNK_CHARS = int(os.getenv("MAX_CHUNK_CHARS", "1500"))
HARD_CHAR_LIMIT = int(os.getenv("HARD_CHAR_LIMIT", "2000"))

# ── Output Token Limits ──
TOKENS_PER_CHAR = int(os.getenv("TOKENS_PER_CHAR", "5"))
OUTPUT_HEADROOM = float(os.getenv("OUTPUT_HEADROOM", "2.0"))
MIN_OUTPUT_TOKENS = 64
MAX_OUTPUT_TOKENS = int(os.getenv("MAX_OUTPUT_TOKENS", "8192"))

API_PROVIDERS = {
    "deepseek": {
        "api_key_env": "DEEPSEEK_API_KEY",
        "base_url": os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
        "model": os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash"),
    },
    "freemodel": {
        "api_key_env": "FREEMODEL_API_KEY",
        "base_url": os.getenv("FREEMODEL_BASE_URL", "https://api.freemodel.dev/v1"),
        "model": os.getenv("FREEMODEL_MODEL", "gpt-4o"),
    },
}
API_PROVIDER_ALIASES = {
    "1": "deepseek",
    "d": "deepseek",
    "deepseek": "deepseek",
    "2": "freemodel",
    "f": "freemodel",
    "free": "freemodel",
    "freemodel": "freemodel",
}

client = None
MODEL_NAME = None
API_PROVIDER_NAME = None


def configure_api_client(provider_name: Optional[str] = None, interactive: bool = True):
    """
    Configure the OpenAI client.
    Uses provider_name if specified; otherwise prompts interactively if running in a TTY,
    or falls back to DEFAULT_API_PROVIDER from .env.
    """
    global client, MODEL_NAME, API_PROVIDER_NAME

    default_provider = DEFAULT_API_PROVIDER
    default_provider = API_PROVIDER_ALIASES.get(default_provider, default_provider)
    if default_provider not in API_PROVIDERS:
        print(
            f"Warning: DEFAULT_API_PROVIDER={default_provider!r} is invalid. "
            "Using deepseek."
        )
        default_provider = "deepseek"

    selected_provider = provider_name
    if not selected_provider and interactive and sys.stdin.isatty():
        print("Choose API provider:")
        print("  1. deepseek")
        print("  2. freemodel")
        try:
            selected = input(
                f"API provider [default: {default_provider}]: "
            ).strip().lower()
        except EOFError:
            selected = ""
        selected_provider = API_PROVIDER_ALIASES.get(selected, selected) if selected else default_provider
        while selected_provider not in API_PROVIDERS:
            print("Invalid provider. Enter 1/deepseek or 2/freemodel.")
            try:
                selected = input(
                    f"API provider [default: {default_provider}]: "
                ).strip().lower()
            except EOFError:
                selected = ""
            selected_provider = API_PROVIDER_ALIASES.get(selected, selected) if selected else default_provider
    elif not selected_provider:
        selected_provider = default_provider
    else:
        selected_provider = API_PROVIDER_ALIASES.get(selected_provider.lower(), selected_provider.lower())

    if selected_provider not in API_PROVIDERS:
        print(f"Error: Unknown provider '{selected_provider}'. Available: {list(API_PROVIDERS.keys())}")
        sys.exit(1)

    provider = API_PROVIDERS[selected_provider]
    api_key_env = provider["api_key_env"]
    api_key = os.getenv(api_key_env)
    if not api_key:
        print(f"Error: {api_key_env} not found in environment or .env file.")
        print(f"Please copy .env.example to .env and set your {api_key_env}.")
        sys.exit(1)

    client = OpenAI(
        api_key=api_key,
        base_url=provider["base_url"],
    )
    MODEL_NAME = provider["model"]
    API_PROVIDER_NAME = selected_provider

# ── Chunk sizing ──
# As per recommendation, chunk by standard characters and natural breaks.
# Keep chunks around 1,000 to 2,000 characters.
MAX_CHUNK_CHARS = 1500
HARD_CHAR_LIMIT = 2000  # Absolute ceiling per chunk

# ── Output token limits ──
TOKENS_PER_CHAR = 5  # Conservative: Devanagari char → tokens
OUTPUT_HEADROOM = 2.0  # Output tokens = input_chars × TOKENS_PER_CHAR × headroom
MIN_OUTPUT_TOKENS = 64  # Even 1-char input gets at least this
MAX_OUTPUT_TOKENS = 8192  # DeepSeek V4 Flash allows up to 8192


# ──────────────────────────────────────────────────────────────────
# 2. SYSTEM PROMPT
# ──────────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a Nepali legal OCR correction specialist. Fix OCR errors and return ONLY the corrected Nepali text.

Fix: wrong Devanagari chars (ण↔न, ष↔स, ब↔व, भ↔म, घ↔प, छ↔इ), broken matras/conjuncts, garbled Unicode, section numbering (दफा/धारा in Devanagari numerals), sub-section numbering, Part/Schedule headers, broken paragraphs (join mid-sentence splits), extra line breaks, watermark noise.

RULES:
- Return ONLY corrected text. No markdown, no comments, no JSON, no tables, no explanations, no commentary.
- DO NOT REPEAT text. Output must be similar length to input.
- PRESERVE ALL NUMBERING: Devanagari numerals like १., २., ३., (१), (२), (क), (ख) MUST remain exactly as-is. Never remove or change article/dhara/section numbers. Always use Nepali Unicode numbering, never English numbering.
- Preserve hierarchy. Keep original when uncertain.
- Preserve Devanagari punctuation. No transliteration.
- No summarizing, no paraphrasing, no translation, no modernization. Preserve original wording exactly.
- If input is only corrupted Nepali words/sentences, return only corrected text without explanations.

HIERARCHY (preserve strictly):
- दफा = number followed by dot, ends with colon. Example: २४. उम्मेदवारको नामावली:
- उपदफा = numbers in brackets, inside दफा. Example: (१), (२), (३)
- खण्ड = Nepali letters in brackets, inside उपदफा. Example: (क), (ख), (ग)
Each दफा, उपदफा, and खण्ड MUST start on a NEW LINE. Maintain this nesting:
  दफा
      उपदफा
          खण्ड

STRUCTURE RULES:
- Preserve and reconstruct numbering hierarchy: १. २. ३., (क) (ख) (ग), (१) (२) (३). If numbering is broken/missing, intelligently restore the proper Nepali sequence.
- Preserve legal headers exactly: भाग, अध्याय, अनुसूची, दफा, धारा, सूची, शीर्षक, नियम, उपनियम, परिच्छेद, उपदफा, उपधारा, स्पष्टीकरण.
- If text is merged into a single paragraph, separate each clause/item properly and reconstruct semantic hierarchy.
- Join artificially broken lines caused by PDF wrapping or OCR.
- Remove only visual artifacts: borders, scan noise, shadows, watermark interference, decorative spacing.
- Never invent missing legal content.

Correct this:"""


# ──────────────────────────────────────────────────────────────────
# 3. TEXT NORMALIZER (STRIP PROBLEMATIC CHARACTERS)
# ──────────────────────────────────────────────────────────────────
def normalize_text(text):
    """
    Strip/replace characters that inflate token count or confuse the model.
    Run BEFORE chunking.
    """
    # Remove ZERO WIDTH NON-JOINER (U+200C) — the #1 culprit
    # These are invisible chars that break BPE tokenization, inflating
    # token counts by 2-3× on affected words.
    text = text.replace("\u200c", "")

    # Remove other zero-width characters
    text = text.replace("\u200b", "")  # Zero Width Space
    text = text.replace("\u200d", "")  # Zero Width Joiner
    text = text.replace("\ufeff", "")  # BOM

    # Normalize smart quotes to standard quotes (OCR artifacts)
    text = text.replace("\u201c", '"')  # Left double quotation
    text = text.replace("\u201d", '"')  # Right double quotation
    text = text.replace("\u2018", "'")  # Left single quotation
    text = text.replace("\u2019", "'")  # Right single quotation

    # Normalize dashes
    text = text.replace("\u2013", "-")  # EN DASH → hyphen
    text = text.replace("\u2014", "-")  # EM DASH → hyphen

    # Remove OCR garbage characters
    text = text.replace("€", "")  # Euro sign (OCR artifact)
    text = text.replace("°", ".")  # Degree sign → period (likely OCR'd period)

    return text


# ──────────────────────────────────────────────────────────────────
# 4. OCR TEXT CLEANER
# ──────────────────────────────────────────────────────────────────
def clean_ocr_text(text):
    # Normalize problematic characters first
    text = normalize_text(text)

    text = re.sub(r"(?i)(www\.[a-z0-9\-]+\.gov\.np|lawcommission\.gov\.np)", "", text)
    cleaned_lines = []

    for line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        if not line:
            continue
        if re.fullmatch(r"[0-9०-९]+", line):
            continue

        has_devanagari = bool(re.search(r"[\u0900-\u097F]", line))
        if not has_devanagari:
            if not re.search(r"(?i)\b(schedule|part|section|article)\b", line):
                if len(line) < 10 and not re.search(r"[a-zA-Z]{3,}", line):
                    continue
        line = re.sub(r"^\|\s*", "", line)
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


# ──────────────────────────────────────────────────────────────────
# 5. CHUNKING LOGIC (STRICT SIZE GUARANTEE)
# ──────────────────────────────────────────────────────────────────
def force_split(text, max_chars):
    """
    Guarantees every returned chunk is ≤ max_chars.
    Split priority: sentence (।) → newline → hard character.
    """
    if len(text) <= max_chars:
        return [text]

    # Split on Devanagari sentence endings (purna viram)
    parts = re.split(r"(?<=।)\s*", text)
    chunks = []
    current = ""

    for part in parts:
        if not part.strip():
            continue
        if len(current) + len(part) > max_chars and current:
            chunks.append(current.strip())
            current = part
        else:
            current += part if not current else " " + part

    if current.strip():
        chunks.append(current.strip())

    # If any chunk is still too large, split on newlines
    pass2 = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            pass2.append(chunk)
        else:
            lines = chunk.split("\n")
            sub = ""
            for line in lines:
                if len(sub) + len(line) + 1 > max_chars and sub:
                    pass2.append(sub.strip())
                    sub = line
                else:
                    sub += "\n" + line if sub else line
            if sub.strip():
                pass2.append(sub.strip())

    # Last resort: hard character split
    final = []
    for chunk in pass2:
        if len(chunk) <= max_chars:
            final.append(chunk)
        else:
            for start in range(0, len(chunk), max_chars):
                piece = chunk[start : start + max_chars].strip()
                if piece:
                    final.append(piece)

    return [c for c in final if c]


def smart_chunk(text, max_chars=MAX_CHUNK_CHARS):
    """
    Two-pass chunking with guaranteed size limit.
    """
    # Split on legal structure markers
    marker_re = (
        r"(?=\n[ \t]*(?:दफा|धारा|भाग|परिच्छेद|अनुसूची|Schedule|Part|Section|Article)\b)"
    )
    parts = [p.strip() for p in re.split(marker_re, text) if p.strip()]

    if len(parts) <= 1:
        parts = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]

    # Merge small parts up to max_chars
    merged = []
    cur = ""
    for p in parts:
        if len(cur) + len(p) > max_chars and len(cur) > 0:
            merged.append(cur.strip())
            cur = p
        else:
            cur += "\n" + p if cur else p

    if cur.strip():
        merged.append(cur.strip())

    # Force-split EVERY chunk that exceeds hard limit
    out = []
    for chunk in merged:
        if len(chunk) > HARD_CHAR_LIMIT:
            out.extend(force_split(chunk, HARD_CHAR_LIMIT))
        else:
            out.append(chunk)

    return out if out else [text]


# ──────────────────────────────────────────────────────────────────
# 6. DOCX I/O
# ──────────────────────────────────────────────────────────────────
def extract_text_from_docx(filepath):
    doc = Document(filepath)
    texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n".join(texts)


# ──────────────────────────────────────────────────────────────────
# 7. OUTPUT VALIDATION
# ──────────────────────────────────────────────────────────────────
def detect_repetition(text):
    """
    Detect CONSECUTIVE repeated blocks — the hallucination pattern.

    Legal Nepali text naturally repeats phrases like 'प्रत्येक नागरिकलाई'
    across different clauses (each clause has different content after the
    repeated opener). Model hallucination, by contrast, repeats the SAME
    full block consecutively (back-to-back).

    Strategy: Split text into lines. Check if any sequence of N consecutive
    lines repeats immediately after itself. This catches the looping pattern
    without flagging natural legal repetition.
    """
    if len(text) < 100:
        return False

    # Method 1: Check for consecutive identical lines
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if len(lines) >= 4:
        # Check if 2+ consecutive lines repeat back-to-back
        for window in [3, 2, 1]:
            for i in range(len(lines) - window * 2 + 1):
                block = lines[i : i + window]
                next_block = lines[i + window : i + window * 2]
                if block == next_block:
                    # Verify it's not just short duplicate lines
                    block_text = " ".join(block)
                    if len(block_text) > 15:
                        return True

    # Method 2: Check for long repeated substring (40+ chars, 4+ times)
    # This catches single-line looping where the model repeats a phrase
    # back-to-back without newlines
    for length in [50, 40]:
        if len(text) < length * 2:
            continue
        for start in range(0, min(len(text) - length, 300), 15):
            pattern = text[start : start + length]
            count = text.count(pattern)
            if count >= 4:
                return True

    # Method 3: Sentence-level consecutive repetition
    # Split on purna viram and check for same sentence back-to-back
    sentences = [s.strip() for s in re.split(r"।", text) if s.strip()]
    if len(sentences) >= 3:
        consecutive = 1
        for i in range(1, len(sentences)):
            if sentences[i] == sentences[i - 1] and len(sentences[i]) > 10:
                consecutive += 1
                if consecutive >= 3:
                    return True
            else:
                consecutive = 1

    return False


def validate_output(input_text, output_text):
    """
    Validate that the model's output is a reasonable correction.
    Returns (is_valid, reason).
    """
    if not output_text or not output_text.strip():
        return False, "empty output"

    input_len = len(input_text)
    output_len = len(output_text)

    # Output should not be drastically larger than input
    # (correction task should produce ~same size text)
    # 4× threshold: corrections can expand abbreviations, fix broken
    # words etc., so we allow generous headroom
    if output_len > input_len * 4 and output_len > 200:
        return (
            False,
            f"output {output_len} chars is {output_len / input_len:.1f}x input {input_len} chars",
        )

    # Check for repetition (model hallucination)
    if detect_repetition(output_text):
        return False, "repetitive output detected"

    return True, "ok"


# ──────────────────────────────────────────────────────────────────
# 8. TOKEN LOGGING & SANITIZATION
# ──────────────────────────────────────────────────────────────────
try:
    tokenizer = tiktoken.get_encoding("cl100k_base")
except Exception:
    tokenizer = None


def get_token_count(text):
    if not text:
        return 0
    if tokenizer:
        try:
            return len(tokenizer.encode(text, disallowed_special=()))
        except Exception:
            return len(text) * 5
    return len(text) * 5


def log_error(error_msg, filename="", context=""):
    try:
        with open(ERROR_LOG_PATH, "a", encoding="utf-8") as f:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"[{timestamp}] ERROR in File: {filename}\n")
            f.write(f"Context: {context}\n")
            f.write(f"{error_msg}\n")
            f.write("-" * 60 + "\n")
    except Exception as e:
        print(f"Failed to write error log: {e}")


def log_culprit(culprit_text, filename, token_count, context):
    try:
        with open(CULPRIT_LOG_PATH, "a", encoding="utf-8") as f:
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"[{timestamp}] CULPRIT SKIPPED\n")
            f.write(f"File: {filename}\n")
            f.write(f"Section/Context: {context[:500]}\n")
            f.write(f"Tokens: {token_count}, Chars: {len(culprit_text)}\n")
            f.write(f"Raw Text (repr): {repr(culprit_text)}\n")
            f.write("-" * 60 + "\n")
    except Exception as e:
        print(f"Failed to write culprit log: {e}")


def sanitize_chunk_from_culprits(text, filename, context_for_logging):
    """
    Two-pass sanitizer:
    Pass 1 (fast): Scan each character individually. If a single char tokenizes
            to an absurd number of tokens (>= 10), it's garbage — drop it.
    Pass 2 (recursive): If the whole chunk still has an abnormal token ratio,
            use divide-and-conquer to isolate the culprit segment.

    cl100k_base is used as a proxy tokenizer. DeepSeek's actual tokenizer is
    worse on Devanagari, so thresholds are set conservatively.
    """
    if not text:
        return text

    # ── Pass 1: Per-character scan for single-char token bombs ──
    cleaned_chars = []
    for ch in text:
        ch_tokens = get_token_count(ch)
        if ch_tokens >= 10:
            # A single character producing 10+ tokens is definitely garbage
            log_culprit(
                ch,
                filename,
                ch_tokens,
                context_for_logging + "\n[Pass1: per-char scan]",
            )
        else:
            cleaned_chars.append(ch)
    text = "".join(cleaned_chars)

    if not text:
        return text

    # ── Pass 2: Recursive divide-and-conquer for multi-char culprits ──
    return _recursive_sanitize(text, filename, context_for_logging)


def _recursive_sanitize(text, filename, context_for_logging):
    """Recursive helper for sanitize_chunk_from_culprits."""
    chars = len(text)
    if chars == 0:
        return text

    tokens = get_token_count(text)
    ratio = tokens / max(chars, 1)

    # Normal Devanagari: 3-5 tok/char. Flag ratio > 8 (lowered from 15
    # because cl100k_base under-estimates vs DeepSeek's tokenizer).
    # Also flag absolute token count > 6000 (well under the 8192 context).
    is_unusual = ratio > 8 or tokens > 6000

    if not is_unusual:
        return text

    # If small enough, we found the exact culprit
    if chars <= 10:
        log_culprit(
            text, filename, tokens, context_for_logging + "\n[Pass2: recursive]"
        )
        return ""  # Skip this bad part

    # Go deep down to find the exact culprit (Divide and Conquer)
    mid = chars // 2
    left = _recursive_sanitize(text[:mid], filename, context_for_logging)
    right = _recursive_sanitize(text[mid:], filename, context_for_logging)

    return left + right


# ──────────────────────────────────────────────────────────────────
# 8.1 DYNAMIC TOKEN CALCULATION
# ──────────────────────────────────────────────────────────────────
def calc_max_tokens(chunk_text):
    """
    Return max output tokens. We use 8192 to avoid truncating Nepali text mid-sentence
    due to Token Bloat.
    """
    return 8192


# ──────────────────────────────────────────────────────────────────
# 9. API CALL (WITH VALIDATION)
# ──────────────────────────────────────────────────────────────────
def call_api(chunk_text, max_tokens, filename="", context=""):
    """
    Single API call. Returns (text, finish_reason, usage) tuple.
    """
    global stop_execution
    if client is None:
        configure_api_client()

    retries = 5
    delay = 5

    for attempt in range(retries):
        if stop_execution:
            return None, "error", None
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": chunk_text},
                ],
                temperature=0.3,
                frequency_penalty=0.8,  # High to prevent repetition
                presence_penalty=0.3,
                top_p=0.8,
                max_tokens=max_tokens,
            )

            choice = response.choices[0]
            text = choice.message.content
            usage = response.usage

            if usage:
                ratio = (
                    f" ratio={usage.completion_tokens / max(usage.prompt_tokens, 1):.1f}x"
                    if usage.prompt_tokens > 0
                    else ""
                )
                print(
                    f"     📊 Tokens: in={usage.prompt_tokens} "
                    f"out={usage.completion_tokens}/{max_tokens}{ratio}"
                )

            return text, choice.finish_reason, usage

        except RateLimitError as e:
            print(f"  ⏳ Rate limit hit! Waiting {delay}s...")
            log_error(f"RateLimitError: {e}", filename, context)
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            print(f"  ❌ Error: {e}")
            log_error(f"API Error: {e}\n{traceback.format_exc()}", filename, context)
            time.sleep(delay)
            delay *= 2

    return None, "error", None


def generate_corrected_text(chunk_text, filename="", chunk_idx=0, depth=0):
    """
    Generate corrected text with output validation.
    """
    global stop_execution
    if stop_execution:
        return ""

    context = (
        f"Chunk index: {chunk_idx}, Depth: {depth}\nExcerpt: {chunk_text[:200]}..."
    )

    # Pre-sanitize to drop token explosion culprits
    sanitized_text = sanitize_chunk_from_culprits(chunk_text, filename, context)
    if not sanitized_text.strip():
        return ""

    max_tokens = calc_max_tokens(sanitized_text)

    text, finish_reason, usage = call_api(sanitized_text, max_tokens, filename, context)

    # ── If the API itself failed (all retries exhausted), don't drill down ──
    # Drilling down on API errors just multiplies failed requests.
    if finish_reason == "error":
        print(
            f"  ❌ API failed for chunk {chunk_idx} at depth {depth}. Using original."
        )
        log_error(
            "API returned error after all retries. Chunk kept as original.",
            filename,
            context,
        )
        return sanitized_text

    is_valid = True
    reason = ""
    if finish_reason == "length":
        is_valid = False
        reason = "Truncated by max_tokens limit"
    elif text and text.strip():
        is_valid, reason = validate_output(sanitized_text, text.strip())
    else:
        is_valid = False
        reason = "Empty output"

    if is_valid and text:
        return text.strip()

    # Drill down logic for invalid outputs (hallucinations, truncations, etc)
    char_count = len(sanitized_text)

    # If the text is very short or we reached max drill-down depth, it's the culprit!
    if char_count <= 10 or depth >= 4:
        print(
            f"  ❌ Culprit identified at depth {depth} ({char_count} chars). "
            f"Reason: {reason}. Skipping."
        )
        log_culprit(
            sanitized_text,
            filename,
            usage.completion_tokens if usage else 0,
            context + "\nReason: " + reason,
        )
        return ""

    print(
        f"  ⚠️  Output rejected ({reason}) at depth={depth} "
        f"({char_count} chars). Splitting and retrying..."
    )

    split_size = max(char_count // 2, 1)
    sub_chunks = force_split(sanitized_text, split_size)

    # If force_split fails to break it down further, hard split it
    if len(sub_chunks) == 1 and sub_chunks[0] == sanitized_text and char_count > 1:
        mid = char_count // 2
        sub_chunks = [sanitized_text[:mid], sanitized_text[mid:]]

    corrected_parts = []
    for i, sub in enumerate(sub_chunks):
        if sub.strip():
            print(f"     ↳ Sub-chunk {i + 1}/{len(sub_chunks)} ({len(sub)} chars)")
            corrected_parts.append(
                generate_corrected_text(sub, filename, chunk_idx, depth + 1)
            )
            if stop_execution:
                break
    return "\n".join(corrected_parts)


# ──────────────────────────────────────────────────────────────────
# 10. POST-PROCESSING: JOIN BROKEN PARAGRAPHS
# ──────────────────────────────────────────────────────────────────
def _is_sentence_end(line):
    """Check if a line ends with a sentence-terminating marker."""
    stripped = line.rstrip()
    if not stripped:
        return True
    # Ends with purna viram, question mark, exclamation, colon, comma,
    # full stop (period), or closing paren/bracket
    if stripped[-1] in "।?!:,.;)\u0964":
        return True
    # Ends with Devanagari danda variants
    if stripped.endswith("।।"):
        return True
    return False


def _is_structural_line(line):
    """Check if a line is a legal structure header that should NOT be joined."""
    stripped = line.strip()
    if not stripped:
        return False
    # Dafa headings like "१३. विदेशी लगानी..." (Devanagari number + dot)
    if re.match(r"^[\d१२३४५६७८९०]+\.\s", stripped):
        return True
    # Upadafa markers like "(२)", "(१)", "(क)", etc.
    if re.match(r"^\([\d१२३४५६७८९०कखगघङचछजझञ]+\)", stripped):
        return True
    # Part/Schedule/Dafa headings
    if re.match(
        r"^(भाग|परिच्छेद|अनुसूची|दफा|धारा|उपदफा|उपधारा|Schedule|Part|Section|Article)\b",
        stripped,
        re.IGNORECASE,
    ):
        return True
    # Khand markers
    if re.match(r"^[\(\(][कखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसह][\)\)]", stripped):
        return True
    # स्पष्टीकरण (Explanation)
    if stripped.startswith("स्पष्टीकरण"):
        return True
    return False


def join_broken_paragraphs(text):
    """
    Join lines that are clearly mid-sentence continuations.
    A line is joined to the previous if:
    - The previous line does NOT end with a sentence-ending marker
      (।, :, comma, full stop, etc.)
    - The current line is NOT a structural heading (dafa, upadafa, etc.)
    - The current line does NOT start with an indented sub-item
    """
    lines = text.split("\n")
    if len(lines) <= 1:
        return text

    result = [lines[0]]
    for i in range(1, len(lines)):
        current = lines[i]
        previous = result[-1]

        # Don't join empty lines (they mark paragraph boundaries)
        if not current.strip():
            result.append(current)
            continue

        # Don't join if current line is a structural element
        if _is_structural_line(current):
            result.append(current)
            continue

        # Join if previous line didn't end a sentence
        if previous.strip() and not _is_sentence_end(previous):
            result[-1] = previous.rstrip() + " " + current.lstrip()
        else:
            result.append(current)

    return "\n".join(result)


# ──────────────────────────────────────────────────────────────────
# 10.1 POST-PROCESSING: FORMAT DAFA / UPADAFA BOUNDARIES
# ──────────────────────────────────────────────────────────────────
def format_dafa_upadafa(text):
    """
    Split running text at dafa, upadafa, and khanda boundaries.

    Dafa pattern:   "...। १३. ..."  (purna biram + Devanagari number + dot)
        → Insert newline before the number so it starts on its own line.

    Upadafa pattern: "...। (२) ..." (purna biram + parenthesized Devanagari number)
        → Insert newline before the paren so it starts on its own indented line.

    Khanda pattern: "...। (क) ..." (purna biram + parenthesized Nepali letter)
        → Insert newline before the paren so it starts on its own indented line.

    The purna biram (।) is the key identifier in all cases.
    """
    # Dafa: "। <number>." → split so the number starts a new line
    # Matches: । followed by optional space, then Devanagari/Arabic digit(s), then dot
    text = re.sub(
        r"।\s*([\d१२३४५६७८९०]+\.\s)",
        r"।\n\1",
        text,
    )

    # Upadafa: "। (<number>)" → split so the paren group starts a new line
    # Matches: । followed by optional space, then (Devanagari/Arabic digit(s))
    text = re.sub(
        r"।\s*(\([\d१२३४५६७८९०]+\))",
        r"।\n\1",
        text,
    )

    # Khanda: "। (क)" → split so the khanda starts a new line
    # Matches: । followed by optional space, then (Nepali letter(s))
    text = re.sub(
        r"।\s*(\([कखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसह]+\))",
        r"।\n\1",
        text,
    )

    return text


# ──────────────────────────────────────────────────────────────────
# 11. DOCX RECONSTRUCTION
# ──────────────────────────────────────────────────────────────────
def build_corrected_docx(text, original_filename, output_folder):
    # Post-process pipeline:
    # 1. Split at dafa/upadafa boundaries
    text = format_dafa_upadafa(text)
    # 2. Join unusual mid-sentence line breaks
    text = join_broken_paragraphs(text)

    doc = Document()
    re_bhag = re.compile(r"^(भाग|परिच्छेद|अनुसूची|Schedule|Part)\b", re.IGNORECASE)
    # Dafa: Devanagari number + dot, e.g. "१३. विदेशी लगानी..."
    re_dafa_numbered = re.compile(r"^[\d१२३४५६७८९०]+\.\s")
    # Named dafa: "दफा ३" or "धारा १२"
    re_dafa_named = re.compile(r"^(दफा|धारा)\s+[\d१२३४५६७८९०]+")
    # Upadafa: "(२)", "(क)", etc.
    re_upadafa = re.compile(r"^\([\d१२३४५६७८९०कखगघङचछजझञ]+\)")
    re_khand = re.compile(r"^[\(\(][कखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसह][\)\)]")

    for line in text.split("\n"):
        trimmed = line.strip()
        if not trimmed:
            continue

        if re_bhag.match(trimmed):
            # Part / Schedule heading → centered heading
            p = doc.add_heading(trimmed, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif re_dafa_numbered.match(trimmed) or re_dafa_named.match(trimmed):
            # Dafa → new paragraph, bold only the title (up to colon)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Inches(0.1)
            # Split at first colon (: or ः) — bold title, normal content
            colon_match = re.search(r"[:\u0903]", trimmed)
            if colon_match:
                colon_idx = colon_match.end()
                title_part = trimmed[:colon_idx]
                content_part = trimmed[colon_idx:].strip()
                bold_run = p.add_run(title_part)
                bold_run.bold = True
                if content_part:
                    p.add_run(" " + content_part)
            else:
                # No colon → bold the entire line (title-only dafa)
                run = p.add_run(trimmed)
                run.bold = True
        elif re_upadafa.match(trimmed):
            # Upadafa → indented paragraph
            p = doc.add_paragraph(trimmed)
            p.paragraph_format.left_indent = Inches(0.3)
        elif re_khand.match(trimmed):
            p = doc.add_paragraph(trimmed)
            p.paragraph_format.left_indent = Inches(0.6)
        else:
            doc.add_paragraph(trimmed)

    os.makedirs(output_folder, exist_ok=True)
    name, ext = os.path.splitext(original_filename)
    out_path = os.path.join(output_folder, f"{name}_corrected{ext}")
    doc.save(out_path)
    return out_path


# ──────────────────────────────────────────────────────────────────
# 11. MAIN PROCESSING LOOP (MULTI-THREADED)
# ──────────────────────────────────────────────────────────────────
def process_directory(
    input_folder: Union[str, Path] = INPUT_DIR_DEFAULT,
    output_folder: Union[str, Path] = OUTPUT_DIR_DEFAULT,
    max_workers: int = MAX_WORKERS_DEFAULT,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
    hard_char_limit: int = HARD_CHAR_LIMIT,
):
    global stop_execution
    input_dir = Path(input_folder)
    if not input_dir.exists():
        print(f"Warning: Input folder '{input_dir}' does not exist.")
        return

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in '{input_dir.resolve()}'.")
        return

    for file_idx, filepath in enumerate(docx_files, 1):
        if stop_execution:
            break

        filename = filepath.name
        print(
            f"\n{'=' * 60}\n[{file_idx}/{len(docx_files)}] Processing: {filename}\n{'=' * 60}"
        )

        raw_text = extract_text_from_docx(str(filepath))
        cleaned_text = clean_ocr_text(raw_text)
        chunks = smart_chunk(cleaned_text, max_chars=max_chunk_chars)

        # Report chunk stats
        chunk_sizes = [len(c) for c in chunks]
        oversized = sum(1 for s in chunk_sizes if s > hard_char_limit)
        print(
            f"  📄 {len(chunks)} chunks. "
            f"min={min(chunk_sizes)}, max={max(chunk_sizes)}, "
            f"avg={sum(chunk_sizes) // len(chunk_sizes)} chars"
        )
        if oversized:
            print(f"  ⚠️  {oversized} chunks exceed hard limit of {hard_char_limit}!")
        print("  🚀 Starting Parallel Processing...\n")

        corrected_chunks_dict = {}

        try:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {
                    executor.submit(generate_corrected_text, chunk, filename, idx): idx
                    for idx, chunk in enumerate(chunks)
                }

                for future in as_completed(futures):
                    idx = futures[future]
                    try:
                        corrected_chunks_dict[idx] = future.result()
                        print(f"  ✅ Completed Chunk {idx + 1}/{len(chunks)}")
                    except Exception as exc:
                        print(f"  ❌ Chunk {idx + 1} exception: {exc}")
                        log_error(
                            f"Executor Error: {exc}\n{traceback.format_exc()}",
                            filename,
                            f"Chunk index {idx}",
                        )
                        corrected_chunks_dict[idx] = chunks[idx]
        except KeyboardInterrupt:
            print("\n🛑 Ctrl+C detected! Stopping execution gracefully...")
            stop_execution = True
            sys.exit(1)

        # ── Retry pass for failed chunks ──
        failed_idxs = [
            idx
            for idx in range(len(chunks))
            if corrected_chunks_dict.get(idx) == chunks[idx]  # still original = failed
        ]
        if failed_idxs and not stop_execution:
            print(f"\n  🔄 Retrying {len(failed_idxs)} failed chunks...")
            for idx in failed_idxs:
                if stop_execution:
                    break
                print(f"     ↳ Retry chunk {idx + 1}/{len(chunks)}")
                try:
                    result = generate_corrected_text(
                        chunks[idx], filename, idx, depth=0
                    )
                    if result and result != chunks[idx]:
                        corrected_chunks_dict[idx] = result
                        print(f"     ✅ Retry succeeded for chunk {idx + 1}")
                    else:
                        print(
                            f"     ⚠️  Retry returned same/empty for chunk {idx + 1}, keeping original"
                        )
                except Exception as exc:
                    print(f"     ❌ Retry failed for chunk {idx + 1}: {exc}")

        # Reconstruct in order
        ordered_chunks = [
            corrected_chunks_dict.get(i, chunks[i]) for i in range(len(chunks))
        ]
        full_corrected_text = "\n\n".join(ordered_chunks)

        out_path = build_corrected_docx(full_corrected_text, filename, output_folder)
        print(f"\n  🎉 Saved: {out_path}\n")


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Nepali Legal OCR Correction Pipeline using OpenAI-compatible LLMs."
    )
    parser.add_argument(
        "-i",
        "--input-dir",
        type=Path,
        default=Path(INPUT_DIR_DEFAULT),
        help=f"Input directory containing .docx documents (default: {INPUT_DIR_DEFAULT} or env INPUT_DIR).",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        default=Path(OUTPUT_DIR_DEFAULT),
        help=f"Output directory to save corrected .docx documents (default: {OUTPUT_DIR_DEFAULT} or env OUTPUT_DIR).",
    )
    parser.add_argument(
        "-p",
        "--provider",
        type=str,
        default=None,
        help="API provider to use ('deepseek' or 'freemodel'). Defaults to interactive prompt or env DEFAULT_API_PROVIDER.",
    )
    parser.add_argument(
        "-w",
        "--workers",
        type=int,
        default=MAX_WORKERS_DEFAULT,
        help=f"Number of parallel worker threads for LLM requests (default: {MAX_WORKERS_DEFAULT} or env MAX_WORKERS).",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=MAX_CHUNK_CHARS,
        help=f"Target character count per semantic chunk (default: {MAX_CHUNK_CHARS} or env MAX_CHUNK_CHARS).",
    )
    parser.add_argument(
        "--hard-limit",
        type=int,
        default=HARD_CHAR_LIMIT,
        help=f"Hard character limit per chunk ceiling (default: {HARD_CHAR_LIMIT} or env HARD_CHAR_LIMIT).",
    )
    parser.add_argument(
        "--non-interactive",
        action="store_true",
        help="Run non-interactively, skipping provider prompt and using default from .env or CLI.",
    )
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    is_interactive = not args.non_interactive and args.provider is None and sys.stdin.isatty()
    configure_api_client(provider_name=args.provider, interactive=is_interactive)

    print("Starting Nepali Legal OCR Correction")
    print(f"  Provider: {API_PROVIDER_NAME} ({MODEL_NAME})")
    print(
        f"  Config: chunk={args.chunk_size}, hard_limit={args.hard_limit}, "
        f"max_out={MAX_OUTPUT_TOKENS}, tok/char={TOKENS_PER_CHAR}, "
        f"headroom={OUTPUT_HEADROOM}x, workers={args.workers}"
    )
    print(f"  Input:    {args.input_dir.resolve()}")
    print(f"  Output:   {args.output_dir.resolve()}")
    print(f"  Logs:     {ERROR_LOG_PATH}")
    print(f"  Culprits: {CULPRIT_LOG_PATH}")

    input_dir = args.input_dir
    output_dir = args.output_dir
    if not input_dir.exists():
        print(f"\nCreating input directory: {input_dir}")
        input_dir.mkdir(parents=True, exist_ok=True)
        print(f"Please place your Nepali legal .docx files in '{input_dir}' and rerun.")
        return 0

    process_directory(
        input_folder=input_dir,
        output_folder=output_dir,
        max_workers=args.workers,
        max_chunk_chars=args.chunk_size,
        hard_char_limit=args.hard_limit,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
